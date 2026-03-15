#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
论文格式修改工具 - Markdown/Word一键转换为格式化Word
支持中国硕士毕业论文格式
"""

import argparse
import json
import os
import subprocess
import sys
import tempfile
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("请先安装依赖: pip install python-docx")
    sys.exit(1)


# 默认中国硕士论文模板
DEFAULT_TEMPLATE = {
    "name": "中国硕士毕业论文",
    "page": {
        "paper_size": "A4",
        "margin": {
            "top": 2.5,
            "bottom": 2.5,
            "left": 3.0,
            "right": 2.0
        }
    },
    "title": {
        "font_name": "黑体",
        "font_size": 22,
        "bold": True,
        "alignment": "center"
    },
    "heading1": {
        "font_name": "黑体",
        "font_size": 16,
        "bold": True,
        "alignment": "center",
        "before_spacing": 24,
        "after_spacing": 18
    },
    "heading2": {
        "font_name": "黑体",
        "font_size": 15,
        "bold": True,
        "alignment": "left",
        "before_spacing": 18,
        "after_spacing": 12
    },
    "heading3": {
        "font_name": "宋体",
        "font_size": 14,
        "bold": True,
        "alignment": "left",
        "before_spacing": 12,
        "after_spacing": 6
    },
    "body": {
        "font_name": "宋体",
        "font_size": 12,
        "line_spacing": 1.5,
        "alignment": "justify"
    },
    "code": {
        "font_name": "Consolas",
        "font_size": 10,
        "border": True,
        "shading": "#F0F0F0"
    },
    "toc": {
        "include": True,
        "title": "目 录",
        "font_name": "黑体",
        "font_size": 16
    }
}


def find_pandoc():
    """查找pandoc路径"""
    paths = [
        "pandoc",
        "/usr/bin/pandoc",
        "/usr/local/bin/pandoc",
        os.path.expanduser("~/.local/bin/pandoc"),
        os.path.expanduser("/home/ubuntu/.local/bin/pandoc")
    ]
    for path in paths:
        try:
            result = subprocess.run([path, "--version"], capture_output=True, timeout=5)
            if result.returncode == 0:
                return path
        except:
            pass
    return None


def word_to_markdown(input_file, output_file=None):
    """将Word文档转换为Markdown"""
    pandoc = find_pandoc()
    if not pandoc:
        print("警告: 未安装pandoc，无法转换Word文件")
        print("请安装: sudo apt install pandoc 或 pip install pandoc")
        return None
        
    if output_file is None:
        output_file = tempfile.mktemp(suffix=".md")
        
    cmd = [pandoc, input_file, "-o", output_file]
    try:
        subprocess.run(cmd, capture_output=True, check=True, timeout=30)
        return output_file
    except subprocess.CalledProcessError as e:
        print(f"Word转换失败: {e}")
        return None
    except Exception as e:
        print(f"转换错误: {e}")
        return None


def extract_format_from_markdown(md_file):
    """从Markdown文件提取格式信息"""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 分析Markdown结构
    format_info = {
        "has_title": False,
        "headings": {"h1": 0, "h2": 0, "h3": 0},
        "has_code": False,
        "has_tables": False,
        "has_lists": False,
        "has_images": False,
    }
    
    lines = content.split('\n')
    for line in lines:
        if line.startswith('# '):
            format_info["has_title"] = True
            format_info["headings"]["h1"] += 1
        elif line.startswith('## '):
            format_info["headings"]["h2"] += 1
        elif line.startswith('### '):
            format_info["headings"]["h3"] += 1
        elif line.startswith('```'):
            format_info["has_code"] = True
        elif '|' in line and '---' in line:
            format_info["has_tables"] = True
        elif line.strip().startswith(('- ', '* ', '+ ')):
            format_info["has_lists"] = True
        elif line.startswith('!['):
            format_info["has_images"] = True
    
    return format_info


class PaperFormatter:
    """论文格式修改器"""
    
    def __init__(self, template=None):
        self.template = template or DEFAULT_TEMPLATE.copy()
        self.doc = None
        self.headings = []
        
    def create_document(self):
        """创建Word文档"""
        self.doc = Document()
        section = self.doc.sections[0]
        margin = self.template.get("page", {}).get("margin", {})
        section.top_margin = Cm(margin.get("top", 2.5))
        section.bottom_margin = Cm(margin.get("bottom", 2.5))
        section.left_margin = Cm(margin.get("left", 3.0))
        section.right_margin = Cm(margin.get("right", 2.0))
        return self.doc
    
    def set_paragraph_format(self, para, format_config):
        """设置段落格式"""
        alignment = format_config.get("alignment", "left")
        if alignment == "center":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == "right":
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == "justify":
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        if "before_spacing" in format_config:
            para.paragraph_format.space_before = Pt(format_config["before_spacing"])
        if "after_spacing" in format_config:
            para.paragraph_format.space_after = Pt(format_config["after_spacing"])
            
        if len(para.runs) == 0:
            para.add_run()
        run = para.runs[0]
        run.font.name = format_config.get("font_name", "宋体")
        run.font.size = Pt(format_config.get("font_size", 12))
        
        if format_config.get("bold"):
            run.font.bold = True
            
        return para
    
    def add_heading(self, text, level=1):
        """添加标题"""
        if level == 1:
            fmt = self.template.get("heading1", {})
        elif level == 2:
            fmt = self.template.get("heading2", {})
        else:
            fmt = self.template.get("heading3", {})
            
        para = self.doc.add_paragraph(text)
        self.set_paragraph_format(para, fmt)
        self.headings.append({"text": text, "level": level, "para": para})
        return para
    
    def add_body_text(self, text):
        """添加正文"""
        fmt = self.template.get("body", {})
        para = self.doc.add_paragraph(text)
        self.set_paragraph_format(para, fmt)
        para.paragraph_format.line_spacing = fmt.get("line_spacing", 1.5)
        return para
    
    def add_code_block(self, text):
        """添加代码块"""
        fmt = self.template.get("code", {})
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = fmt.get("font_name", "Consolas")
        run.font.size = Pt(fmt.get("font_size", 10))
        
        try:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), fmt.get("shading", "#F0F0F0").replace("#", ""))
            para._p.get_or_add_pPr().append(shading_elm)
        except:
            pass
            
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        return para
    
    def add_list(self, text, ordered=False):
        """添加列表"""
        fmt = self.template.get("body", {})
        lines = text.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if line.startswith(('- ', '* ', '+ ')):
                line = line[2:]
            elif line and line[0].isdigit() and '. ' in line:
                line = line[line.index('. ')+2:]
                
            para = self.doc.add_paragraph(line)
            para.style = 'List Bullet' if not ordered else 'List Number'
            
            if len(para.runs) == 0:
                para.add_run()
            run = para.runs[0]
            run.font.name = fmt.get("font_name", "宋体")
            run.font.size = Pt(fmt.get("font_size", 12))
            
        return para
    
    def add_table(self, markdown_table):
        """添加表格"""
        lines = markdown_table.strip().split('\n')
        if len(lines) < 3:
            return None
            
        rows = []
        for line in lines:
            if line.strip().startswith('|') and '---' not in line:
                cells = [c.strip() for c in line.strip().split('|')[1:-1]]
                rows.append(cells)
                
        if not rows:
            return None
            
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if len(paragraph.runs) == 0:
                        paragraph.add_run()
                    run = paragraph.runs[0]
                    run.font.name = "宋体"
                    run.font.size = Pt(10.5)
        return table
    
    def parse_markdown(self, md_content):
        """解析Markdown内容并转换为Word"""
        lines = md_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            if not line.strip():
                i += 1
                continue
                
            # 标题处理
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                
                if level == 1:
                    if len(self.headings) == 0:
                        fmt = self.template.get("title", {})
                        para = self.doc.add_paragraph(text)
                        self.set_paragraph_format(para, fmt)
                    else:
                        self.add_heading(text, 1)
                elif level == 2:
                    self.add_heading(text, 2)
                elif level >= 3:
                    self.add_heading(text, 3)
                    
            # 代码块
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                self.add_code_block('\n'.join(code_lines))
                
            # 表格
            elif '|' in line:
                if i + 1 < len(lines) and '---' in lines[i + 1]:
                    table_lines = [line]
                    i += 1
                    while i < len(lines) and '|' in lines[i]:
                        table_lines.append(lines[i])
                        i += 1
                    self.add_table('\n'.join(table_lines))
                    i -= 1
                    
            # 列表
            stripped = line.strip()
            is_bullet = stripped.startswith(('- ', '* ', '+ '))
            is_numbered = stripped and stripped[0].isdigit() and '. ' in stripped[:5]
            if is_bullet or is_numbered:
                list_text = line
                while i + 1 < len(lines):
                    next_line = lines[i+1].strip()
                    if not next_line:
                        break
                    if next_line.startswith(('- ', '* ', '+ ')) or (next_line and next_line[0].isdigit() and '. ' in next_line[:5]):
                        i += 1
                        list_text += '\n' + lines[i]
                    else:
                        break
                self.add_list(list_text)
                
            # 普通段落
            else:
                para_text = line
                while i + 1 < len(lines):
                    next_line = lines[i+1].strip()
                    if not next_line:
                        break
                    if next_line.startswith(('#', '-', '*', '+', '```', '|', '![', '>')):
                        break
                    if next_line and next_line[0].isdigit() and '. ' in next_line[:5]:
                        break
                    i += 1
                    para_text += '\n' + lines[i]
                self.add_body_text(para_text)
                
            i += 1
    
    def generate_toc(self):
        """生成目录"""
        toc_config = self.template.get("toc", {})
        if not toc_config.get("include", False):
            return
            
        toc_title = toc_config.get("title", "目 录")
        fmt = toc_config.copy()
        para = self.doc.add_paragraph(toc_title)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if len(para.runs) == 0:
            para.add_run()
        run = para.runs[0]
        run.font.name = fmt.get("font_name", "黑体")
        run.font.size = Pt(fmt.get("font_size", 14))
        
        for heading in self.headings:
            para = self.doc.add_paragraph()
            run = para.add_run(heading["text"])
            run.font.name = "宋体"
            run.font.size = Pt(12)
            
        self.doc.add_page_break()
    
    def convert(self, input_file, output_file):
        """转换Markdown到Word"""
        input_ext = Path(input_file).suffix.lower()
        md_content = ""
        
        if input_ext in ['.docx', '.doc']:
            # Word输入：先转为Markdown
            temp_md = word_to_markdown(input_file)
            if temp_md:
                with open(temp_md, 'r', encoding='utf-8') as f:
                    md_content = f.read()
                os.unlink(temp_md)
            else:
                print("Word转换失败，尝试直接读取...")
                return False
        else:
            # Markdown输入：直接读取
            with open(input_file, 'r', encoding='utf-8') as f:
                md_content = f.read()
        
        # 创建文档
        self.create_document()
        
        # 解析Markdown
        self.parse_markdown(md_content)
        
        # 生成目录
        self.generate_toc()
        
        # 保存文档
        self.doc.save(output_file)
        return True


def load_template(template_path):
    """加载模板文件"""
    if not template_path or template_path == "default":
        return DEFAULT_TEMPLATE.copy()
        
    with open(template_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def list_templates():
    """列出可用模板"""
    print("可用模板:")
    print("  default - 中国硕士毕业论文格式（默认）")
    
    template_dir = Path(__file__).parent / "templates"
    if template_dir.exists():
        for f in template_dir.glob("*.json"):
            print(f"  {f.stem} - {f.name}")


def main():
    parser = argparse.ArgumentParser(
        description="论文格式修改工具 - Markdown/Word一键转换为格式化Word"
    )
    parser.add_argument("input", nargs="?", help="输入文件 (.md 或 .docx)")
    parser.add_argument("output", nargs="?", help="输出文件 (.docx)")
    parser.add_argument("-t", "--template", help="模板文件路径或模板名")
    parser.add_argument("--list-templates", action="store_true", help="列出可用模板")
    parser.add_argument("--extract", action="store_true", help="从Markdown提取格式生成模板")
    
    args = parser.parse_args()
    
    if args.list_templates:
        list_templates()
        return
    
    if args.extract and args.input:
        info = extract_format_from_markdown(args.input)
        print("检测到的格式:")
        print(f"  标题: H1={info['headings']['h1']}, H2={info['headings']['h2']}, H3={info['headings']['h3']}")
        print(f"  代码块: {info['has_code']}")
        print(f"  表格: {info['has_tables']}")
        print(f"  列表: {info['has_lists']}")
        print(f"  图片: {info['has_images']}")
        return
        
    if not args.input or not args.output:
        parser.print_help()
        print("\n示例:")
        print("  python paper_formatter.py 论文.md 论文.docx")
        print("  python paper_formatter.py 论文.docx 论文.docx  # 支持Word输入!")
        print("  python paper_formatter.py 论文.md 论文.docx --template custom.json")
        return
        
    # 加载模板
    template = load_template(args.template)
    
    # 执行转换
    formatter = PaperFormatter(template)
    
    try:
        formatter.convert(args.input, args.output)
        print(f"✓ 转换成功: {args.output}")
        print(f"  模板: {template.get('name', '自定义')}")
    except Exception as e:
        import traceback
        print(f"✗ 转换失败: {e}")
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
