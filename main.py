import os
import json
from datetime import datetime
from kivy.app import App
from kivy.uix.screenmanager import ScreenManager, Screen
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.gridlayout import GridLayout
from kivy.uix.scrollview import ScrollView
from kivy.uix.label import Label
from kivy.uix.button import Button
from kivy.uix.textinput import TextInput
from kivy.uix.spinner import Spinner
from kivy.uix.switch import Switch
from kivy.uix.popup import Popup
from kivy.core.window import Window
from kivy.properties import ObjectProperty, StringProperty
from kivy.lang import Builder
from plyer import filechooser
import shutil
import tempfile

Window.softinput_mode = "resize"


class HomeScreen(Screen):
    pass


class ConvertScreen(Screen):
    input_file = StringProperty("")
    output_file = StringProperty("")


class SettingsScreen(Screen):
    pass


class PaperFormatterApp(App):
    title = "论文格式转换器"
    
    def build(self):
        # Load template
        self.template = self.load_template()
        
        sm = ScreenManager()
        sm.add_widget(HomeScreen(name='home'))
        sm.add_widget(ConvertScreen(name='convert'))
        sm.add_widget(SettingsScreen(name='settings'))
        return sm
    
    def load_template(self):
        template_path = os.path.join(os.path.dirname(__file__), 
                                     'templates', 'china_master.json')
        if os.path.exists(template_path):
            with open(template_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        return self.get_default_template()
    
    def get_default_template(self):
        return {
            "name": "中国硕士论文",
            "page": {"paper_size": "A4", "margin": {"top": 2.5, "bottom": 2.5, "left": 3.0, "right": 2.0}},
            "title": {"font_name": "黑体", "font_size": 22, "bold": True, "alignment": "center"},
            "heading1": {"font_name": "黑体", "font_size": 16, "bold": True, "alignment": "center"},
            "heading2": {"font_name": "黑体", "font_size": 15, "bold": True, "alignment": "left"},
            "heading3": {"font_name": "宋体", "font_size": 12, "bold": False, "alignment": "left"},
            "body": {"font_name": "宋体", "font_size": 12, "line_spacing": 1.5, "alignment": "justify"},
            "code": {"font_name": "Consolas", "font_size": 10, "border": True, "shading": "#F0F0F0"},
            "toc": {"include": True, "title": "目 录"}
        }
    
    def choose_input_file(self):
        try:
            path = filechooser.open_file(title="选择输入文件",
                                         filters=[("文档", "*.md *.docx"), ("所有文件", "*.*")])
            if path:
                return path[0]
        except:
            # Fallback for testing
            return ""
        return ""
    
    def convert_document(self, input_path, output_path, template=None):
        """Convert document using python-docx"""
        from docx import Document
        from docx.shared import Pt, Cm, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        import markdown
        
        if template is None:
            template = self.template
        
        # Create document
        doc = Document()
        
        # Set page margins
        section = doc.sections[0]
        margin = template.get('page', {}).get('margin', {})
        section.top_margin = Cm(margin.get('top', 2.5))
        section.bottom_margin = Cm(margin.get('bottom', 2.5))
        section.left_margin = Cm(margin.get('left', 3.0))
        section.right_margin = Cm(margin.get('right', 2.0))
        
        # Read input
        if input_path.endswith('.md'):
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            html = markdown.markdown(content)
            # Simple conversion - in production use proper markdown parser
            lines = content.split('\n')
            in_code_block = False
            code_content = []
            
            for line in lines:
                if line.startswith('```'):
                    in_code_block = not in_code_block
                    if not in_code_block:
                        # Add code paragraph
                        self.add_paragraph(doc, '\n'.join(code_content), 'code', template)
                        code_content = []
                    continue
                
                if in_code_block:
                    code_content.append(line)
                elif line.startswith('# '):
                    self.add_paragraph(doc, line[2:], 'title', template)
                elif line.startswith('## '):
                    self.add_paragraph(doc, line[3:], 'heading1', template)
                elif line.startswith('### '):
                    self.add_paragraph(doc, line[4:], 'heading2', template)
                elif line.strip():
                    self.add_paragraph(doc, line, 'body', template)
        
        elif input_path.endswith('.docx'):
            # Copy and apply formatting
            src_doc = Document(input_path)
            for para in src_doc.paragraphs:
                new_para = doc.add_paragraph()
                new_para.text = para.text
                self.apply_formatting(new_para, 'body', template)
        
        # Save
        doc.save(output_path)
        return output_path
    
    def add_paragraph(self, doc, text, style, template):
        style_cfg = template.get(style, template.get('body', {}))
        para = doc.add_paragraph(text)
        self.apply_formatting(para, style, template)
        return para
    
    def apply_formatting(self, paragraph, style, template):
        style_cfg = template.get(style, template.get('body', {}))
        
        # Font
        font = paragraph.style.font
        font.name = style_cfg.get('font_name', '宋体')
        font.size = Pt(style_cfg.get('font_size', 12))
        font.bold = style_cfg.get('bold', False)
        
        # Chinese font
        try:
            paragraph.style.font.name = style_cfg.get('font_name', '宋体')
            paragraph._element.rPr.rFonts.set(qn('w:eastAsia'), style_cfg.get('font_name', '宋体'))
        except:
            pass
        
        # Alignment
        align = style_cfg.get('alignment', 'left')
        if align == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align == 'justify':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Spacing
        before = style_cfg.get('before_spacing', 0)
        after = style_cfg.get('after_spacing', 0)
        paragraph.space_before = Pt(before)
        paragraph.space_after = Pt(after)
        
        # Line spacing
        line_spacing = style_cfg.get('line_spacing', 1.5)
        paragraph.paragraph_format.line_spacing = line_spacing


if __name__ == '__main__':
    PaperFormatterApp().run()
